"""
Motor de transcripción de AureaTranscribe.

Utiliza faster-whisper para transcripción local con timestamps,
soporte multi-idioma simultáneo (detección por segmento),
y opcionalmente pyannote.audio para diarización de hablantes.
FFmpeg se encarga de la conversión de formatos.
"""

import os
import tempfile
import subprocess
import json
import logging
from datetime import datetime
from pathlib import Path
from dataclasses import dataclass, field, asdict
from typing import List, Optional, Callable, Dict

# ── Parche PyTorch 2.6+: permitir carga de modelos pyannote/speechbrain ──
# PyTorch 2.6 cambió weights_only=True por defecto en torch.load,
# lo que impide cargar los checkpoints de pyannote y speechbrain.
try:
    import torch
    import torch.serialization

    # Método 1: Registrar globals seguros conocidos
    try:
        torch.serialization.add_safe_globals([torch.torch_version.TorchVersion])
    except Exception:
        pass

    # Método 2: Parchar torch.load Y torch.serialization.load
    _original_torch_load = torch.serialization.load if hasattr(torch.serialization, 'load') else torch.load

    def _patched_torch_load(*args, **kwargs):
        kwargs["weights_only"] = False
        return _original_torch_load(*args, **kwargs)

    torch.load = _patched_torch_load
    if hasattr(torch.serialization, 'load'):
        torch.serialization.load = _patched_torch_load

except ImportError:
    pass

logger = logging.getLogger("AureaTranscribe")

LANGUAGE_NAMES = {
    "es": "Español", "en": "Inglés", "fr": "Francés", "de": "Alemán",
    "it": "Italiano", "pt": "Portugués", "ca": "Catalán", "gl": "Gallego",
    "eu": "Euskera", "ar": "Árabe", "zh": "Chino", "ja": "Japonés",
    "ko": "Coreano", "ru": "Ruso", "nl": "Neerlandés", "pl": "Polaco",
    "ro": "Rumano", "uk": "Ucraniano", "sv": "Sueco", "da": "Danés",
    "no": "Noruego", "fi": "Finés", "el": "Griego", "cs": "Checo",
    "hu": "Húngaro", "tr": "Turco", "he": "Hebreo", "hi": "Hindi",
}


# ─────────────────────────────────────────────────────────────
# Modelos de datos
# ─────────────────────────────────────────────────────────────

@dataclass
class Segment:
    """Un segmento de transcripción."""
    start: float
    end: float
    text: str
    speaker: Optional[str] = None
    language: Optional[str] = None  # Idioma detectado en este segmento

@dataclass
class TranscriptionResult:
    """Resultado completo de una transcripción."""
    segments: List[Segment] = field(default_factory=list)
    full_text: str = ""
    language: str = ""
    languages_detected: List[str] = field(default_factory=list)
    language_probability: float = 0.0
    duration: float = 0.0
    num_speakers: int = 0
    speaker_names: Dict[str, str] = field(default_factory=dict)
    original_filename: str = ""
    transcription_date: str = ""
    transcription_time: str = ""

    def to_dict(self):
        return {
            "segments": [asdict(s) for s in self.segments],
            "full_text": self.full_text,
            "language": self.language,
            "languages_detected": self.languages_detected,
            "language_probability": self.language_probability,
            "duration": self.duration,
            "num_speakers": self.num_speakers,
            "speaker_names": self.speaker_names,
            "original_filename": self.original_filename,
            "transcription_date": self.transcription_date,
            "transcription_time": self.transcription_time,
        }


# ─────────────────────────────────────────────────────────────
# Utilidades de audio
# ─────────────────────────────────────────────────────────────

def get_ffmpeg_path() -> str:
    """Localiza FFmpeg en el sistema."""
    import shutil
    import sys
    # En ejecutable PyInstaller, buscar en la carpeta de extraccion
    if getattr(sys, '_MEIPASS', None):
                exe = 'ffmpeg.exe' if os.name == 'nt' else 'ffmpeg'
                bundled = os.path.join(sys._MEIPASS, exe)
                if os.path.exists(bundled):
                                return bundled
    path = shutil.which("ffmpeg")
    if path:
        return path
    raise FileNotFoundError(
        "FFmpeg no encontrado. Instálalo desde https://ffmpeg.org/download.html"
    )


def get_ffprobe_path() -> str:
    """Localiza FFprobe en el sistema."""
    import shutil
    import sys
    # En ejecutable PyInstaller, buscar en la carpeta de extraccion
    if getattr(sys, '_MEIPASS', None):
        exe = 'ffprobe.exe' if os.name == 'nt' else 'ffprobe'
        bundled = os.path.join(sys._MEIPASS, exe)
        if os.path.exists(bundled):
            return bundled
    path = shutil.which("ffprobe")
    if path:
        return path
    # Fallback: derivar de ffmpeg
    ffmpeg = get_ffmpeg_path()
    probe = ffmpeg.replace("ffmpeg", "ffprobe")
    if os.path.exists(probe):
        return probe
    raise FileNotFoundError("FFprobe no encontrado.")


def get_audio_duration(file_path: str) -> float:
    """Obtiene la duración de un archivo de audio/vídeo en segundos."""
    try:
        result = subprocess.run(
            [get_ffprobe_path(), "-v", "quiet", "-print_format", "json",
             "-show_format", file_path],
            capture_output=True, text=True, timeout=30
        )
        info = json.loads(result.stdout)
        return float(info["format"]["duration"])
    except Exception:
        return 0.0


def convert_to_wav(input_path: str, output_path: str) -> str:
    """Convierte cualquier formato de audio/vídeo a WAV 16kHz mono."""
    ffmpeg = get_ffmpeg_path()
    cmd = [
        ffmpeg, "-y", "-i", input_path,
        "-ar", "16000", "-ac", "1", "-c:a", "pcm_s16le",
        output_path,
    ]
    result = subprocess.run(cmd, capture_output=True, text=True, timeout=600)
    if result.returncode != 0:
        raise RuntimeError(f"Error al convertir audio: {result.stderr[:500]}")
    return output_path


# ─────────────────────────────────────────────────────────────
# Motor de transcripción
# ─────────────────────────────────────────────────────────────

class TranscriptionEngine:
    """Motor principal de transcripción con Whisper y diarización opcional."""

    AVAILABLE_MODELS = {
        "tiny":     "Muy rápido, calidad básica (~1 GB RAM)",
        "base":     "Rápido, calidad aceptable (~1 GB RAM)",
        "small":    "Equilibrado, buena calidad (~2 GB RAM)",
        "medium":   "Lento, alta calidad (~5 GB RAM)",
        "large-v3": "Muy lento, máxima calidad (~10 GB RAM)",
    }

    def __init__(self):
        self.model = None
        self.model_name = None
        self.diarization_pipeline = None
        self._diarization_available = None

    def load_model(self, model_name: str = "medium",
                   progress_callback: Optional[Callable] = None):
        """Carga el modelo Whisper."""
        from faster_whisper import WhisperModel

        if progress_callback:
            progress_callback("loading_model", f"Cargando modelo '{model_name}'...")

        device = "cpu"
        compute_type = "int8"

        try:
            import torch
            if torch.cuda.is_available():
                device = "cuda"
                compute_type = "float16"
        except ImportError:
            pass

        self.model = WhisperModel(
            model_name, device=device, compute_type=compute_type,
            download_root=os.path.join(
                os.path.expanduser("~"), ".cache", "aurea_transcribe", "models"
            ),
        )
        self.model_name = model_name

        if progress_callback:
            progress_callback("model_loaded",
                              f"Modelo '{model_name}' cargado en {device.upper()}")
        logger.info(f"Modelo {model_name} cargado en {device} ({compute_type})")

    def check_diarization(self) -> bool:
        if self._diarization_available is not None:
            return self._diarization_available
        try:
            from pyannote.audio import Pipeline
            self._diarization_available = True
        except ImportError:
            self._diarization_available = False
        return self._diarization_available

    @staticmethod
    def _patch_huggingface_hub():
        """Parche de compatibilidad: pyannote antiguo pasa use_auth_token
        a huggingface_hub nuevo que ya solo acepta token."""
        try:
            import huggingface_hub as _hfh

            for fn_name in ("hf_hub_download", "snapshot_download",
                            "model_info", "list_repo_files"):
                original = getattr(_hfh, fn_name, None)
                if original is None or getattr(original, "_aurea_patched", False):
                    continue

                def _make_wrapper(orig):
                    def wrapper(*args, **kwargs):
                        if "use_auth_token" in kwargs:
                            kwargs["token"] = kwargs.pop("use_auth_token")
                        return orig(*args, **kwargs)
                    wrapper._aurea_patched = True
                    wrapper.__name__ = orig.__name__
                    return wrapper

                setattr(_hfh, fn_name, _make_wrapper(original))
        except Exception:
            pass

    def load_diarization(self, hf_token: Optional[str] = None,
                         progress_callback: Optional[Callable] = None):
        if not self.check_diarization():
            raise ImportError(
                "Para usar diarización, instala: pip install pyannote.audio\n"
                "Y obtén un token en https://huggingface.co/settings/tokens"
            )

        # Aplicar parche de compatibilidad ANTES de importar pyannote
        self._patch_huggingface_hub()

        from pyannote.audio import Pipeline

        if progress_callback:
            progress_callback("loading_diarization", "Cargando modelo de diarización...")

        token = hf_token or os.environ.get("HF_TOKEN")
        if not token:
            raise ValueError(
                "Se necesita un token de HuggingFace para la diarización.\n"
                "Obtén uno en https://huggingface.co/settings/tokens"
            )
        # Inyectar token vía variable de entorno + parámetro directo
        os.environ["HF_TOKEN"] = token
        os.environ["HUGGING_FACE_HUB_TOKEN"] = token
        self.diarization_pipeline = Pipeline.from_pretrained(
            "pyannote/speaker-diarization-3.1",
            use_auth_token=token,
        )

        # Nota: MPS (Apple Silicon GPU) no es compatible con speechbrain,
        # por lo que la diarización se ejecuta en CPU.
        # En un M4 nativo arm64, el rendimiento en CPU es suficiente.
        print("[AureaTranscribe] ✓ Modelo de diarización cargado (CPU arm64)")

        if progress_callback:
            progress_callback("diarization_loaded", "Modelo de diarización cargado")

    def transcribe(
        self,
        file_path: str,
        language: Optional[str] = None,
        multilingual: bool = False,
        diarize: bool = False,
        hf_token: Optional[str] = None,
        progress_callback: Optional[Callable] = None,
    ) -> TranscriptionResult:
        """
        Transcribe un archivo de audio o vídeo.

        Args:
            file_path: Ruta al archivo.
            language: Código de idioma (ej: 'es'). None = autodetección.
            multilingual: Si True, detecta el idioma de cada segmento individualmente.
            diarize: Activar diarización de hablantes.
            hf_token: Token de HuggingFace (solo para diarización).
            progress_callback: Función callback(stage, message, pct).
        """
        if self.model is None:
            self.load_model(progress_callback=progress_callback)

        now = datetime.now()

        # ── Paso 1: Convertir a WAV si es necesario ──
        temp_wav = None
        audio_path = file_path

        ext = Path(file_path).suffix.lower()
        if ext != ".wav":
            if progress_callback:
                progress_callback("converting", "Convirtiendo formato de audio...")
            temp_wav = tempfile.NamedTemporaryFile(suffix=".wav", delete=False)
            temp_wav.close()
            audio_path = convert_to_wav(file_path, temp_wav.name)

        try:
            duration = get_audio_duration(audio_path)

            # ── Paso 2: Transcripción ──
            if progress_callback:
                progress_callback("transcribing", "Transcribiendo audio...")

            # Si es multilingüe, no fijamos idioma para que Whisper
            # detecte automáticamente por segmento
            transcribe_lang = None if multilingual else language

            segments_raw, info = self.model.transcribe(
                audio_path,
                language=transcribe_lang,
                beam_size=5,
                best_of=5,
                patience=1.0,
                condition_on_previous_text=True,
                vad_filter=True,
                vad_parameters=dict(
                    min_silence_duration_ms=500,
                    speech_pad_ms=400,
                ),
                word_timestamps=True,
            )

            segments_list = []
            detected_languages = set()

            for seg in segments_raw:
                seg_lang = getattr(seg, "language", None) or info.language
                detected_languages.add(seg_lang)

                segments_list.append(Segment(
                    start=round(seg.start, 2),
                    end=round(seg.end, 2),
                    text=seg.text.strip(),
                    language=seg_lang,
                ))
                if progress_callback and duration > 0:
                    pct = min(int((seg.end / duration) * 80), 80)
                    progress_callback("transcribing_progress",
                                      f"Transcribiendo... {pct}%", pct)

            # ── Paso 3: Diarización (opcional) ──
            if diarize and self.check_diarization():
                if self.diarization_pipeline is None:
                    self.load_diarization(hf_token, progress_callback)
                if progress_callback:
                    progress_callback("diarizing", "Identificando hablantes...", 85)
                import time as _time
                print(f"[AureaTranscribe] Iniciando diarización del audio ({round(duration)}s)...")
                _t0 = _time.time()
                diarization = self.diarization_pipeline(audio_path)
                _elapsed = round(_time.time() - _t0, 1)
                print(f"[AureaTranscribe] ✓ Diarización completada en {_elapsed}s")
                segments_list = self._assign_speakers(segments_list, diarization)

            # ── Paso 4: Construir resultado ──
            if progress_callback:
                progress_callback("finalizing", "Finalizando transcripción...", 95)

            speakers = set(s.speaker for s in segments_list if s.speaker)
            # Crear mapa de nombres por defecto
            speaker_names = {}
            for sp in sorted(speakers):
                speaker_names[sp] = sp  # "Hablante 1" → "Hablante 1"

            result = TranscriptionResult(
                segments=segments_list,
                full_text=self._build_full_text(segments_list),
                language=info.language,
                languages_detected=sorted(detected_languages),
                language_probability=round(info.language_probability, 3),
                duration=round(duration, 2),
                num_speakers=len(speakers),
                speaker_names=speaker_names,
                original_filename=Path(file_path).name,
                transcription_date=now.strftime("%d/%m/%Y"),
                transcription_time=now.strftime("%H:%M:%S"),
            )

            if progress_callback:
                progress_callback("done", "Transcripción completada", 100)

            return result

        finally:
            if temp_wav and os.path.exists(temp_wav.name):
                os.unlink(temp_wav.name)

    def _assign_speakers(self, segments: List[Segment], diarization) -> List[Segment]:
        for segment in segments:
            best_speaker = None
            best_overlap = 0
            for turn, _, speaker in diarization.itertracks(yield_label=True):
                overlap_start = max(segment.start, turn.start)
                overlap_end = min(segment.end, turn.end)
                overlap = max(0, overlap_end - overlap_start)
                if overlap > best_overlap:
                    best_overlap = overlap
                    best_speaker = speaker
            if best_speaker:
                num = int(best_speaker.split("_")[-1]) + 1 if "_" in best_speaker else 1
                segment.speaker = f"Hablante {num}"
        return segments

    def _build_full_text(self, segments: List[Segment]) -> str:
        lines = []
        current_speaker = None
        for seg in segments:
            if seg.speaker and seg.speaker != current_speaker:
                current_speaker = seg.speaker
                lines.append(f"\n[{current_speaker}]")
            lines.append(seg.text)
        return " ".join(lines).strip()


# ─────────────────────────────────────────────────────────────
# Metadatos de la aplicación y perfil del usuario
# ─────────────────────────────────────────────────────────────

APP_INFO = {
    "app_name": "AureaTranscribe",
    "app_version": "1.0",
    "brand": "Áurea Laboral",
}

# Perfil del usuario: se carga desde la configuración persistente.
# Si el usuario no ha rellenado sus datos, se deja vacío.
DEFAULT_USER_PROFILE = {
    "name": "",
    "firm": "",
    "email": "",
}


# ─────────────────────────────────────────────────────────────
# Exportadores
# ─────────────────────────────────────────────────────────────

def format_timestamp(seconds: float, srt_format: bool = False) -> str:
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs = int(seconds % 60)
    millis = int((seconds % 1) * 1000)
    sep = "," if srt_format else "."
    return f"{hours:02d}:{minutes:02d}:{secs:02d}{sep}{millis:03d}"


def format_duration_human(seconds: float) -> str:
    h = int(seconds // 3600)
    m = int((seconds % 3600) // 60)
    s = int(seconds % 60)
    if h > 0:
        return f"{h}h {m}m {s}s"
    if m > 0:
        return f"{m}m {s}s"
    return f"{s}s"


def _get_languages_str(result: TranscriptionResult) -> str:
    langs = result.languages_detected or [result.language]
    return ", ".join(LANGUAGE_NAMES.get(l, l) for l in langs)


def _resolve_speaker(speaker: Optional[str],
                     speaker_names: Dict[str, str]) -> Optional[str]:
    """Resuelve el nombre real del hablante usando el mapa de nombres."""
    if not speaker:
        return None
    return speaker_names.get(speaker, speaker)


# ── TXT ──

def export_txt(result: TranscriptionResult, output_path: str,
               include_timestamps: bool = True, user_profile: dict = None):
    profile = user_profile or DEFAULT_USER_PROFILE
    with open(output_path, "w", encoding="utf-8") as f:
        f.write("TRANSCRIPCIÓN\n")
        f.write(f"Fecha: {result.transcription_date}  Hora: {result.transcription_time}\n")
        if profile.get("name"):
            f.write(f"Profesional: {profile['name']}\n")
        if profile.get("firm"):
            f.write(f"Empresa: {profile['firm']}\n")
        if profile.get("email"):
            f.write(f"Contacto: {profile['email']}\n")
        f.write(f"Archivo fuente: {result.original_filename}\n")
        f.write(f"Idioma(s): {_get_languages_str(result)}\n")
        f.write(f"Duración: {format_duration_human(result.duration)}\n")
        if result.num_speakers > 0:
            f.write(f"Interlocutores ({result.num_speakers}):")
            for key, name in result.speaker_names.items():
                f.write(f" {name};")
            f.write("\n")
        f.write("=" * 60 + "\n\n")

        current_speaker = None
        for seg in result.segments:
            resolved = _resolve_speaker(seg.speaker, result.speaker_names)
            if resolved and resolved != current_speaker:
                current_speaker = resolved
                f.write(f"\n--- {current_speaker} ---\n")
            if include_timestamps:
                ts = format_timestamp(seg.start)
                f.write(f"[{ts}] {seg.text}\n")
            else:
                f.write(f"{seg.text}\n")

        f.write(f"\n{'=' * 60}\n")
        f.write(f"Powered by {APP_INFO['app_name']} · {APP_INFO['brand']}\n")


# ── Markdown ──

def export_md(result: TranscriptionResult, output_path: str,
              user_profile: dict = None):
    profile = user_profile or DEFAULT_USER_PROFILE
    with open(output_path, "w", encoding="utf-8") as f:
        f.write("# Transcripción\n\n")
        f.write(f"**Fecha:** {result.transcription_date} · {result.transcription_time}  \n")
        if profile.get("name"):
            f.write(f"**Profesional:** {profile['name']}  \n")
        if profile.get("firm"):
            f.write(f"**Empresa:** {profile['firm']}  \n")
        if profile.get("email"):
            f.write(f"**Contacto:** {profile['email']}  \n")
        f.write(f"**Archivo fuente:** {result.original_filename}  \n")
        f.write(f"**Idioma(s):** {_get_languages_str(result)}  \n")
        f.write(f"**Duración:** {format_duration_human(result.duration)}  \n")
        if result.num_speakers > 0:
            names = ", ".join(result.speaker_names.values())
            f.write(f"**Interlocutores ({result.num_speakers}):** {names}  \n")
        f.write("\n---\n\n")

        current_speaker = None
        for seg in result.segments:
            resolved = _resolve_speaker(seg.speaker, result.speaker_names)
            if resolved and resolved != current_speaker:
                current_speaker = resolved
                f.write(f"\n### {current_speaker}\n\n")
            ts = format_timestamp(seg.start)
            f.write(f"`{ts}` {seg.text}  \n")

        f.write(f"\n---\n*Powered by {APP_INFO['app_name']} · "
                f"{APP_INFO['brand']}*\n")


# ── SRT ──

def export_srt(result: TranscriptionResult, output_path: str):
    with open(output_path, "w", encoding="utf-8") as f:
        for i, seg in enumerate(result.segments, 1):
            start = format_timestamp(seg.start, srt_format=True)
            end = format_timestamp(seg.end, srt_format=True)
            resolved = _resolve_speaker(seg.speaker, result.speaker_names)
            text = f"[{resolved}] {seg.text}" if resolved else seg.text
            f.write(f"{i}\n{start} --> {end}\n{text}\n\n")


# ── VTT ──

def export_vtt(result: TranscriptionResult, output_path: str):
    with open(output_path, "w", encoding="utf-8") as f:
        f.write("WEBVTT\n\n")
        for i, seg in enumerate(result.segments, 1):
            start = format_timestamp(seg.start)
            end = format_timestamp(seg.end)
            resolved = _resolve_speaker(seg.speaker, result.speaker_names)
            text = f"[{resolved}] {seg.text}" if resolved else seg.text
            f.write(f"{i}\n{start} --> {end}\n{text}\n\n")


# ── DOCX ──

def export_docx(result: TranscriptionResult, output_path: str,
                user_profile: dict = None):
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    profile = user_profile or DEFAULT_USER_PROFILE

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Aptos"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)

    # ── Cabecera ──
    title = doc.add_heading("TRANSCRIPCIÓN", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(48, 48, 48)

    doc.add_paragraph("")

    # ── Tabla de metadatos ──
    meta_data = [
        ("Fecha y hora", f"{result.transcription_date} · {result.transcription_time}"),
    ]
    if profile.get("name"):
        meta_data.append(("Profesional", profile["name"]))
    if profile.get("firm"):
        meta_data.append(("Empresa", profile["firm"]))
    if profile.get("email"):
        meta_data.append(("Contacto", profile["email"]))
    meta_data.extend([
        ("Archivo fuente", result.original_filename),
        ("Idioma(s)", _get_languages_str(result)),
        ("Duración", format_duration_human(result.duration)),
    ])
    if result.num_speakers > 0:
        names = ", ".join(result.speaker_names.values())
        meta_data.append(("Interlocutores", f"{result.num_speakers} — {names}"))

    table = doc.add_table(rows=len(meta_data), cols=2)
    table.style = "Table Grid"
    for i, (label, value) in enumerate(meta_data):
        cell_label = table.cell(i, 0)
        cell_value = table.cell(i, 1)
        run_l = cell_label.paragraphs[0].add_run(label)
        run_l.bold = True
        run_l.font.size = Pt(10)
        run_l.font.color.rgb = RGBColor(80, 70, 50)
        run_v = cell_value.paragraphs[0].add_run(value)
        run_v.font.size = Pt(10)
        # Ancho de columna
        cell_label.width = Cm(4)

    doc.add_paragraph("")

    # ── Separador ──
    sep = doc.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sep_run = sep.add_run("— CONTENIDO DE LA TRANSCRIPCIÓN —")
    sep_run.font.size = Pt(9)
    sep_run.font.color.rgb = RGBColor(160, 140, 80)
    sep_run.bold = True

    doc.add_paragraph("")

    # ── Contenido ──
    current_speaker = None
    for seg in result.segments:
        resolved = _resolve_speaker(seg.speaker, result.speaker_names)
        if resolved and resolved != current_speaker:
            current_speaker = resolved
            sp_para = doc.add_paragraph()
            sp_para.paragraph_format.space_before = Pt(12)
            sp_run = sp_para.add_run(current_speaker.upper())
            sp_run.bold = True
            sp_run.font.size = Pt(11)
            sp_run.font.color.rgb = RGBColor(160, 140, 80)

        para = doc.add_paragraph()
        ts_run = para.add_run(f"[{format_timestamp(seg.start)}]  ")
        ts_run.font.size = Pt(8)
        ts_run.font.color.rgb = RGBColor(140, 130, 100)
        text_run = para.add_run(seg.text)
        text_run.font.size = Pt(11)

    # ── Pie ──
    doc.add_paragraph("")
    doc.add_paragraph("")
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run(
        f"Powered by {APP_INFO['app_name']} · {APP_INFO['brand']}"
    )
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(160, 140, 80)
    fr.italic = True

    doc.save(output_path)


# ── PDF ──

def export_pdf(result: TranscriptionResult, output_path: str,
               user_profile: dict = None):
    """Exporta a PDF profesional con metadatos completos."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm, mm
    from reportlab.lib.colors import HexColor
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        HRFlowable
    )
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY

    # Colores
    GOLD = HexColor("#c8a55a")
    GOLD_DARK = HexColor("#a08545")
    DARK = HexColor("#1a1a1a")
    GRAY = HexColor("#666666")
    LIGHT_GRAY = HexColor("#f5f3ee")
    WHITE = HexColor("#ffffff")

    profile = user_profile or DEFAULT_USER_PROFILE

    doc = SimpleDocTemplate(
        output_path, pagesize=A4,
        topMargin=2*cm, bottomMargin=2*cm,
        leftMargin=2.5*cm, rightMargin=2.5*cm,
        title="Transcripción — AureaTranscribe",
        author=profile.get("name", ""),
    )

    styles = getSampleStyleSheet()

    # Estilos personalizados
    title_style = ParagraphStyle(
        "AureaTitle", parent=styles["Title"],
        fontName="Helvetica-Bold", fontSize=20,
        textColor=DARK, alignment=TA_CENTER,
        spaceAfter=6,
    )
    subtitle_style = ParagraphStyle(
        "AureaSubtitle", parent=styles["Normal"],
        fontName="Helvetica", fontSize=9,
        textColor=GOLD_DARK, alignment=TA_CENTER,
        spaceAfter=16,
    )
    meta_label_style = ParagraphStyle(
        "MetaLabel", parent=styles["Normal"],
        fontName="Helvetica-Bold", fontSize=9,
        textColor=GOLD_DARK,
    )
    meta_value_style = ParagraphStyle(
        "MetaValue", parent=styles["Normal"],
        fontName="Helvetica", fontSize=9,
        textColor=DARK,
    )
    speaker_style = ParagraphStyle(
        "Speaker", parent=styles["Normal"],
        fontName="Helvetica-Bold", fontSize=10,
        textColor=GOLD_DARK, spaceBefore=14, spaceAfter=4,
    )
    segment_style = ParagraphStyle(
        "Segment", parent=styles["Normal"],
        fontName="Helvetica", fontSize=10,
        textColor=DARK, alignment=TA_JUSTIFY,
        leading=14, spaceAfter=3,
    )
    timestamp_style_inline = (
        '<font face="Courier" size="7" color="#a08545">[{}]</font>&nbsp;&nbsp;'
    )
    footer_style = ParagraphStyle(
        "Footer", parent=styles["Normal"],
        fontName="Helvetica-Oblique", fontSize=7,
        textColor=GOLD_DARK, alignment=TA_CENTER,
        spaceBefore=20,
    )

    elements = []

    # ── Título ──
    elements.append(Paragraph("TRANSCRIPCIÓN", title_style))
    elements.append(Paragraph(
        f"{APP_INFO['app_name']} · {APP_INFO['brand']}",
        subtitle_style
    ))

    # ── Línea decorativa ──
    elements.append(HRFlowable(
        width="100%", thickness=1, color=GOLD, spaceAfter=16
    ))

    # ── Tabla de metadatos ──
    meta_rows = [
        ("Fecha y hora", f"{result.transcription_date} · {result.transcription_time}"),
    ]
    if profile.get("name"):
        meta_rows.append(("Profesional", profile["name"]))
    if profile.get("firm"):
        meta_rows.append(("Empresa", profile["firm"]))
    if profile.get("email"):
        meta_rows.append(("Contacto", profile["email"]))
    meta_rows.extend([
        ("Archivo fuente", result.original_filename),
        ("Idioma(s)", _get_languages_str(result)),
        ("Duración", format_duration_human(result.duration)),
    ])
    if result.num_speakers > 0:
        names = ", ".join(result.speaker_names.values())
        meta_rows.append(("Interlocutores", f"{result.num_speakers} — {names}"))

    meta_table_data = []
    for label, value in meta_rows:
        meta_table_data.append([
            Paragraph(label, meta_label_style),
            Paragraph(value, meta_value_style),
        ])

    meta_table = Table(meta_table_data, colWidths=[4*cm, None])
    meta_table.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ("LINEBELOW", (0, 0), (-1, -2), 0.3, HexColor("#e0ddd5")),
        ("LINEBELOW", (0, -1), (-1, -1), 0.3, HexColor("#e0ddd5")),
    ]))
    elements.append(meta_table)
    elements.append(Spacer(1, 16))

    # ── Separador ──
    elements.append(HRFlowable(
        width="100%", thickness=0.5, color=GOLD, spaceAfter=12
    ))

    # ── Contenido ──
    current_speaker = None
    for seg in result.segments:
        resolved = _resolve_speaker(seg.speaker, result.speaker_names)
        if resolved and resolved != current_speaker:
            current_speaker = resolved
            elements.append(Paragraph(current_speaker.upper(), speaker_style))

        ts_html = timestamp_style_inline.format(format_timestamp(seg.start))
        # Escapar caracteres especiales para XML
        safe_text = (seg.text
                     .replace("&", "&amp;")
                     .replace("<", "&lt;")
                     .replace(">", "&gt;"))
        elements.append(Paragraph(ts_html + safe_text, segment_style))

    # ── Pie ──
    elements.append(Spacer(1, 24))
    elements.append(HRFlowable(width="100%", thickness=0.5, color=GOLD, spaceAfter=8))
    elements.append(Paragraph(
        f"Powered by {APP_INFO['app_name']} v{APP_INFO['app_version']} · "
        f"{APP_INFO['brand']}",
        footer_style
    ))

    doc.build(elements)


# ── JSON ──

def export_json(result: TranscriptionResult, output_path: str):
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(result.to_dict(), f, ensure_ascii=False, indent=2)
